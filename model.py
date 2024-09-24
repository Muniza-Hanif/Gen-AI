import openai
import pytesseract
import pdfplumber
import os
import speech_recognition as sr
from docx import Document
from fpdf import FPDF
from PIL import Image, ImageDraw
import time
from googletrans import Translator

openai.api_key = 'YOUR_API_KEY'  # Replace with your actual API key

def get_gpt_response(prompt):
    "Get a response from OpenAI GPT with retry logic for rate limits."
    for _ in range(5):
        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant."},
                    {"role": "user", "content": prompt}
                ]
            )
            return response['choices'][0]['message']['content']
        except openai.error.RateLimitError:
            print("Rate limit exceeded, retrying in 5 seconds...")
            time.sleep(5)
    return "Sorry, I'm currently unable to process your request."

def recognize_speech():
    """Recognize speech and convert it into text."""
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        print("Listening...")
        audio = recognizer.listen(source)
        try:
            return recognizer.recognize_google(audio)
        except sr.UnknownValueError:
            return "Sorry, I did not understand the speech."
        except sr.RequestError:
            return "Could not request results; check your internet connection."

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file."""
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    return text

def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file."""
    doc = Document(docx_path)
    return '\n'.join([para.text for para in doc.paragraphs])

def extract_text_from_image(image_path):
    """Extract text from an image using OCR."""
    image = Image.open(image_path)
    return pytesseract.image_to_string(image)

def extract_text_from_txt(txt_path):
    """Extract text from a TXT file."""
    with open(txt_path, 'r', encoding='utf-8') as file:
        return file.read()

def translate_text(text, target_language):
    """Translate text to the target language using Google Translate."""
    translator = Translator()
    try:
        if not text.strip():
            return "No text to translate."
        translated = translator.translate(text, dest=target_language)
        return translated.text
    except Exception as e:
        print(f"Translation error: {e}")  # Log the error for debugging
        return "Translation error occurred."

def translate_file(file_path, file_content, file_type, target_language):
    """Translate the content of a file and return a downloadable version."""
    translated_content = translate_text(file_content, target_language)

    if file_type == "pdf":
        output_pdf = os.path.splitext(file_path)[0] + "_translated.pdf"
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for line in translated_content.split("\n"):
            pdf.cell(200, 10, txt=line, ln=True)
        pdf.output(output_pdf)
        return output_pdf

    elif file_type == "docx":
        output_docx = os.path.splitext(file_path)[0] + "_translated.docx"
        doc = Document()
        for para in translated_content.split("\n"):
            doc.add_paragraph(para)
        doc.save(output_docx)
        return output_docx

    elif file_type in ["jpg", "jpeg", "png"]:
        output_image = os.path.splitext(file_path)[0] + "_translated.png"
        img = Image.new('RGB', (500, 300), color=(73, 109, 137))
        d = ImageDraw.Draw(img)
        d.text((10, 10), translated_content, fill=(255, 255, 0))
        img.save(output_image)
        return output_image

    elif file_type == "txt":
        output_txt = os.path.splitext(file_path)[0] + "_translated.txt"
        with open(output_txt, 'w', encoding='utf-8') as file:
            file.write(translated_content)
        return output_txt

def query_content(file_content, user_query):
    """Query the content of a file based on the user's query."""
    prompt = f"Based on the document, answer the following question: {user_query}. Document: {file_content}"
    return get_gpt_response(prompt)

# Example usage:
# file_path = 'your_file.txt'  # Change this to your file path
# file_type = 'txt'
# target_language = 'hi'  # Hindi
# file_content = extract_text_from_txt(file_path)
# translated_file = translate_file(file_path, file_content, file_type, target_language)
# print(f'Translated file saved as: {translated_file}')

